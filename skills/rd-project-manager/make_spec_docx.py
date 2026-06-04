"""
make_spec_docx.py — Generate a styled SPEC.docx from a SPEC.txt source.

Usage:
    python make_spec_docx.py --project-dir <path> --project-name <name>

Example:
    python make_spec_docx.py --project-dir "C:\\MyProject" --project-name "MyProject"

Requirements:
    pip install python-docx
"""
import os
import re
import argparse
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate a polished SPEC.docx from SPEC.txt"
    )
    parser.add_argument(
        "--project-dir", required=True,
        help="Path to the project folder containing <ProjectName>_SPEC.txt"
    )
    parser.add_argument(
        "--project-name", required=True,
        help="Project name (must match the SPEC filename prefix)"
    )
    parser.add_argument(
        "--output", default=None,
        help="Optional: override output DOCX path"
    )
    return parser.parse_args()


# ── Color palette ──────────────────────────────────────────────────────────────
BRAND_DARK   = (31,  55, 100)
BRAND_MID    = (46, 117, 182)
ACCENT_GREY  = (89,  89,  89)
WHITE        = (255, 255, 255)
LIGHT_BLUE   = (235, 243, 251)

STATUS_COLORS = {
    'Planned':        (108, 117, 125),
    'In Progress':    ( 13, 110, 253),
    'In Development': ( 13, 110, 253),
    'Done':           (  0, 128,   0),
    'Released':       (  0, 128,   0),
    'Deprecated':     (180,   0,   0),
}


# ── XML helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def rf(run, name='微軟正黑體', east='微軟正黑體', size=10,
       bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    r2 = OxmlElement('w:rFonts')
    r2.set(qn('w:eastAsia'), east)
    r2.set(qn('w:ascii'), name)
    rPr.insert(0, r2)


def margins(doc, top=2.0, bottom=1.8, left=2.5, right=2.5):
    for sec in doc.sections:
        sec.top_margin    = Cm(top)
        sec.bottom_margin = Cm(bottom)
        sec.left_margin   = Cm(left)
        sec.right_margin  = Cm(right)


# ── Document building blocks ───────────────────────────────────────────────────
def add_hdr_footer(doc, project_name, version):
    for section in doc.sections:
        hdr = section.header
        hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run(f'{project_name}  |  Product Specification  |  v{version}')
        rf(hr, size=8, italic=True, color=ACCENT_GREY)

        ftr = section.footer
        fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run('Confidential — For Internal Use Only    |    ')
        rf(fr, size=8, italic=True, color=ACCENT_GREY)
        fldChar1  = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.text = ' PAGE \\* MERGEFORMAT '
        fldChar2  = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run_el = OxmlElement('w:r')
        run_el.append(fldChar1); run_el.append(instrText); run_el.append(fldChar2)
        fp._p.append(run_el)


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    set_para_bg(p, '1F3764')
    r = p.add_run(f'  {text}')
    rf(r, size=13, bold=True, color=WHITE)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    rf(r, size=11, bold=True, color=BRAND_MID)


def body(doc, text, indent=0, color=None, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    rf(r, size=size, color=color, bold=bold, italic=italic)


def kv(doc, label, value, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{label:<14}')
    rf(r1, size=10, bold=True, color=BRAND_MID)
    r2 = p.add_run(value)
    rf(r2, size=10)


def sep(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run('─' * 90)
    rf(r, size=7, color=(200, 200, 200))


def styled_table(doc, headers, rows, col_widths=None,
                 hdr_fill='2E6DB4', alt_fill='EBF3FB'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], hdr_fill)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        rf(run, size=9, bold=True, color=WHITE)
    for ri, row_data in enumerate(rows):
        cells = tbl.rows[ri + 1].cells
        if ri % 2 == 1:
            for c in cells:
                set_cell_bg(c, alt_fill)
        for ci, val in enumerate(row_data):
            p = cells[ci].paragraphs[0]
            if headers[ci] == 'Status':
                col = STATUS_COLORS.get(str(val), ACCENT_GREY)
                run = p.add_run(str(val))
                rf(run, size=9, bold=True, color=col)
            else:
                run = p.add_run(str(val))
                rf(run, size=9)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()
    return tbl


def get_field(key, text):
    m = re.search(rf'^{key}\s*:\s*(.+)$', text, re.MULTILINE)
    return m.group(1).strip() if m else ''


# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    args = parse_args()

    project_dir  = args.project_dir
    project_name = args.project_name
    txt_path     = os.path.join(project_dir, f"{project_name}_SPEC.txt")
    out_path     = args.output or os.path.join(project_dir, f"{project_name}_SPEC.docx")

    if not os.path.exists(txt_path):
        raise FileNotFoundError(f"SPEC.txt not found: {txt_path}")

    with open(txt_path, encoding="utf-8") as f:
        raw = f.read()

    version = get_field('Version', raw)
    status  = get_field('Status',  raw)
    owner   = get_field('Owner',   raw)
    created = get_field('Created', raw)
    updated = get_field('Updated', raw)

    doc = Document()
    margins(doc)

    # ── Cover page ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_bg(p, '1F3764')
    r = p.add_run(f'  {project_name}  ')
    rf(r, size=28, bold=True, color=WHITE)

    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Product Specification')
    rf(r, size=16, bold=True, color=BRAND_DARK)

    doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'Version {version}    ')
    rf(r1, size=12, color=BRAND_DARK)
    status_col = STATUS_COLORS.get(status, ACCENT_GREY)
    r2 = p.add_run(f'[ {status} ]')
    rf(r2, size=12, bold=True, color=status_col)

    doc.add_paragraph()
    sep(doc)

    styled_table(doc,
        ['Field', 'Value'],
        [
            ['Project',  project_name],
            ['Version',  version],
            ['Status',   status],
            ['Owner',    owner],
            ['Created',  created],
            ['Updated',  updated],
        ],
        col_widths=[4, 11],
        hdr_fill='2E6DB4',
    )

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— CONFIDENTIAL —')
    rf(r, size=9, italic=True, color=(150, 150, 150))

    doc.add_page_break()
    add_hdr_footer(doc, project_name, version)

    # ── 1. Overview ─────────────────────────────────────────────────────────────
    h1(doc, '1.  Overview')
    sep(doc)
    for label, placeholder in [
        ('Purpose',  get_field('Purpose',  raw) or '（描述此專案的目的與解決的問題）'),
        ('Scope',    get_field('Scope',    raw) or '（說明功能範圍，包含哪些、不包含哪些）'),
        ('Platform', get_field('Platform', raw) or '（適用作業系統 / 硬體平台）'),
    ]:
        kv(doc, label, placeholder)

    # ── 2. Features ─────────────────────────────────────────────────────────────
    h1(doc, '2.  Features')
    sep(doc)

    # Parse feature table from TXT
    feat_block = re.search(r'\[2\. FEATURES\].*?\n(.*?)(?=\[3\.|\Z)', raw, re.DOTALL)
    feat_rows = []
    if feat_block:
        for line in feat_block.group(1).splitlines():
            m = re.match(r'\s*(F-\d+)\s*\|\s*(.+?)\s*\|\s*(\S+)\s*\|\s*(.+?)\s*\|\s*(.+)', line)
            if m:
                feat_rows.append([m.group(1), m.group(2), m.group(3), m.group(4), m.group(5)])

    if not feat_rows:
        feat_rows = [['F-01', '（功能名稱）', '0.1.0', 'Planned', '（功能描述）']]

    styled_table(doc,
        ['ID', 'Feature Name', 'Version', 'Status', 'Description'],
        feat_rows,
        col_widths=[1.2, 3.8, 1.8, 2.5, 5.7]
    )

    # ── 3. Architecture ─────────────────────────────────────────────────────────
    h1(doc, '3.  Architecture')
    sep(doc)
    arch_block = re.search(r'\[3\. ARCHITECTURE\](.*?)(?=\[4\.|\Z)', raw, re.DOTALL)
    arch_text = arch_block.group(1).strip() if arch_block else '（描述系統架構、模組劃分、資料流）'
    body(doc, arch_text, indent=0.5, italic=True, color=(120, 120, 120))

    # ── 4. Configuration ────────────────────────────────────────────────────────
    h1(doc, '4.  Configuration')
    sep(doc)
    cfg_block = re.search(r'\[4\. CONFIGURATION\].*?\n(.*?)(?=\[5\.|\Z)', raw, re.DOTALL)
    cfg_rows = []
    if cfg_block:
        for line in cfg_block.group(1).splitlines():
            m = re.match(r'\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+)', line)
            if m and 'Parameter' not in m.group(1) and '─' not in line:
                cfg_rows.append([m.group(1), m.group(2), m.group(3), m.group(4)])
    if not cfg_rows:
        cfg_rows = [['（param_name）', '（value）', '（range）', '（說明）']]
    styled_table(doc,
        ['Parameter', 'Default', 'Range / Values', 'Description'],
        cfg_rows,
        col_widths=[3.5, 2.5, 3.0, 6.0]
    )

    # ── 5. Dependencies ─────────────────────────────────────────────────────────
    h1(doc, '5.  Dependencies')
    sep(doc)
    dep_block = re.search(r'\[5\. DEPENDENCIES\].*?\n(.*?)(?=\[6\.|\Z)', raw, re.DOTALL)
    dep_rows = []
    if dep_block:
        for line in dep_block.group(1).splitlines():
            m = re.match(r'\s*(.+?)\s*\|\s*(.+?)\s*\|\s*(.+)', line)
            if m and 'Name' not in m.group(1) and '─' not in line:
                dep_rows.append([m.group(1), m.group(2), m.group(3)])
    if not dep_rows:
        dep_rows = [['（library/tool）', '（ver）', '（用途）']]
    styled_table(doc,
        ['Name', 'Version', 'Purpose'],
        dep_rows,
        col_widths=[4.0, 2.5, 8.5]
    )

    # ── 6. Change Log ───────────────────────────────────────────────────────────
    h1(doc, '6.  Change Log')
    sep(doc)
    body(doc, '※ 最新版本列於最上方', indent=0.5, italic=True, color=ACCENT_GREY, size=9)
    doc.add_paragraph()

    log_block = re.search(r'\[6\. CHANGE LOG\].*?\n(.*?)(?=={10,}|\Z)', raw, re.DOTALL)
    log_rows = []
    if log_block:
        for line in log_block.group(1).splitlines():
            m = re.match(r'\s*(\d+\.\d+\.\d+)\s*\|\s*(\S+)\s*\|\s*(.+?)\s*\|\s*(.+)', line)
            if m:
                log_rows.append([m.group(1), m.group(2), m.group(3), m.group(4)])
    if not log_rows:
        log_rows = [['0.1.0', '2026-06-04', 'Author', 'Initial draft']]
    styled_table(doc,
        ['Version', 'Date', 'Author', 'Summary'],
        log_rows,
        col_widths=[1.8, 3.0, 3.5, 6.7],
        hdr_fill='1F3764'
    )

    # ── Footer note ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    sep(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f'Generated from {os.path.basename(txt_path)}  |  '
        f'{datetime.date.today()}  |  Auto-generated by rd-project-manager skill'
    )
    rf(r, size=8, italic=True, color=(170, 170, 170))

    doc.save(out_path)
    print(f'✓ Saved: {out_path}')
    print(f'   Source TXT : {txt_path}')
    print(f'   Output DOCX: {out_path}')
    print(f'\n✅ SPEC Word document generated from TXT source.')


if __name__ == '__main__':
    main()
